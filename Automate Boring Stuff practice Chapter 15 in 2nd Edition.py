import PyPDF2
import docx
import Read_Docx

# p. 357

# pdf_file_obj = open('meetingminutes.pdf', 'rb')
# pdf_reader = PyPDF2_.PdfFileReader(pdf_file_obj)
# print(pdf_reader.numPages)
#
# page_obj = pdf_reader.getPage(8)
# print(page_obj.extractText())
# pdf_file_obj.close()

# p. 358

# pdf_reader = PyPDF2_.PdfFileReader(open('encrypted.pdf', 'rb'))
# print(pdf_reader.isEncrypted)
#
# print(pdf_reader.decrypt('rosebud'))
#
# page_obj = pdf_reader.getPage(8)
#
# print(page_obj)

# p. 359

# pdf1_file = open('meetingminutes.pdf', 'rb')
# pdf2_file = open('meetingminutes2.pdf', 'rb')
# pdf1_reader = PyPDF2.PdfFileReader(pdf1_file)
# pdf2_reader = PyPDF2.PdfFileReader(pdf2_file)
# pdf_writer = PyPDF2.PdfFileWriter()
#
# for page_num in range(pdf1_reader.numPages):
#     page_obj = pdf1_reader.getPage(page_num)
#     pdf_writer.addPage(page_obj)
#
# for page_num in range(pdf2_reader.numPages):
#     page_obj = pdf2_reader.getPage(page_num)
#     pdf_writer.addPage(page_obj)
#
# pdf_output_file = open('combinedminutes.pdf', 'wb')
# pdf_writer.write(pdf_output_file)
# pdf_output_file.close()
# pdf1_file.close()
# pdf2_file.close()

# p. 359-60

# minutes_file = open('meetingminutes.pdf', 'rb')
# pdf_reader = PyPDF2.PdfFileReader(minutes_file)
# page = pdf_reader.getPage(0)
# print(page.rotateClockwise(90))
#
# pdf_writer = PyPDF2.PdfFileWriter()
# pdf_writer.addPage(page)
# result_pdf_file = open('rotatedPage.pdf', 'wb')
# pdf_writer.write(result_pdf_file)
# result_pdf_file.close()
# minutes_file.close()

# p. 360-361
# minutes_file = open('meetingminutes.pdf', 'rb')
# pdf_reader = PyPDF2.PdfFileReader(minutes_file)
# minutes_first_page = pdf_reader.getPage(0)
# pdf_watermark_reader = PyPDF2.PdfFileReader(open('watermark.pdf', 'rb'))
# minutes_first_page.mergePage(pdf_watermark_reader.getPage(0))
# pdf_writer = PyPDF2.PdfFileWriter()
# pdf_writer.addPage(minutes_first_page)
#
# for page_num in range(1, pdf_reader.numPages):
#     page_obj = pdf_reader.getPage(page_num)
#     pdf_writer.addPage(page_obj)
#
# result_pdf_file = open('watermarked_cover.pdf', 'wb')
# pdf_writer.write(result_pdf_file)
# minutes_file.close()
# result_pdf_file.close()
# print('all done yoooo')

# p. 361

# pdf_file = open('meetingminutes.pdf', 'rb')
# pdf_reader = PyPDF2.PdfFileReader(pdf_file)
# pdf_writer = PyPDF2.PdfFileWriter()
# for page_num in range(pdf_reader.numPages):
#     pdf_writer.addPage(pdf_reader.getPage(page_num))
#
# pdf_writer.encrypt('swordfish')
# result_pdf = open('encryptedminutes.pdf', 'wb')
# pdf_writer.write(result_pdf)
# result_pdf.close()
#
# print('all done yoooo')

# p. 365

doc = docx.Document('demo.docx')
# print(len(doc.paragraphs))
#
# print(doc.paragraphs[0].text)
#
# print(doc.paragraphs[1].text)
#
# print(len(doc.paragraphs[1].runs))
#
# print(doc.paragraphs[1].runs[0].text)
#
# print(doc.paragraphs[1].runs[2].text)
#
# print(doc.paragraphs[1].runs[3].text)
#
# print(doc.paragraphs[1].runs[4].text)

# p. 366

# print(Read_Docx.getText('demo.docx'))

# You can also adjust getText() to modify the string before returning it. For example, to indent each
# paragraph, replace the append() call in readDocx.py with this:

# fullText.append(' ' + para.text)

# To add a double space between paragraphs, change the join() call code from Read_Docx to this:

# return '\n\n'.join(fullText)

# p. 368

# doc = docx.Document('demo.docx')
# print(doc.paragraphs[0].text)
# print(doc.paragraphs[0].style)  # The exact id may be different:
# doc.paragraphs[0].style = 'Normal'
# print(doc.paragraphs[1].text)
# print((doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.
#        paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text))
#
# doc.paragraphs[1].runs[0].style = 'Quote Char'
# doc.paragraphs[1].runs[1].underline = True
# doc.paragraphs[1].runs[3].underline = True
# doc.save('restyled.docx')

# p. 369

# doc = docx.Document()
# print(doc.add_paragraph('Hello, world!'))
# doc.save('helloworld.docx')
#
# paraObj1 = doc.add_paragraph('This is a second paragraph.')
# paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
# print(paraObj1.add_run(' This text is being added to the second paragraph.'))
# doc.save('multiple paragraphs.docx')

# p. 370

# doc = docx.Document()
# print(doc.add_heading('Header 0', 0))
#
# doc.add_heading('Header 1', 1)
#
# doc.add_heading('Header 2', 2)
#
# doc.add_heading('Header 3', 3)
#
# doc.add_heading('Header 4', 4)
#
# doc.save('headings.docx')

# p. 371

doc = docx.Document()
print(doc.add_paragraph('This is on the first page!'))

doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page, yooooooo!')

doc.add_picture('zophie.png', width=docx.shared.Inches(1),
                height=docx.shared.Cm(4))

doc.save('Two Paaaaage!.docx')
